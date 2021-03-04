
from docx.shared import Inches, Pt



def dodaj_tabele(document, queryset):
    table = document.add_table(rows=0, cols=7)

    lp = 1
    for i in queryset:

        # formatowanie daty
        data = 'w dn. '
        #jesli miesiace sa te same
        if str(i.data_powrotu)[5:7] == str(i.data_wyjazdu)[5:7]:
            #jesli PJ dluzsza niz 1 dzień
            if str(i.data_wyjazdu)[8:10] != str(i.data_powrotu)[8:10]:
                # dzien-
                data += str(i.data_wyjazdu)[8:10] + ' - '
        else:
            #dzien.miesiac-
            data += str(i.data_wyjazdu)[8:10] + '.' + str(i.data_wyjazdu)[5:7] + ' - '
        #dzien.miesiac.rok
        data += str(i.data_powrotu)[8:10] + '.' + str(i.data_powrotu)[5:7] + '.' + str(i.data_powrotu)[0:4] + ' r.'
        # wpisywanie danych do tabeli
        komorki = table.add_row().cells
        komorki[0].text = str(lp) + ')'
        komorki[1].text = i.stopien
        komorki[2].text = i.imie
        komorki[3].text = str(i.nazwisko).upper()
        komorki[4].text = data
        komorki[5].text = 'do m.'
        komorki[6].text = i.miasto
        lp += 1
    return table

# switcher do odpowiedniego sortowania stopni na rozkazie
def switch_stopien(x):
    switcher = {
        'szer. pchor.' : '0',
        'st. szer. pchor.' : '1',
        'kpr. pchor.' : '2',
        'st. kpr. pchor.' : '3',
        'plut. pchor.' : '4',
        'sierż. pchor.' : '5'
    }
    return switcher.get(x, '9')


# switch case do liter w kolejnych podpunktach
def switch_litery(x):
    switcher = {
        1 : 'a) ',
        2: 'b) ',
        3: 'c) ',
        4: 'd) ',
    }
    return switcher.get(x, '9')
# dodanie naglowku dokumentu i
def dodaj_naglowek(document):
    p = document.add_paragraph()
    #p.paragraph_format.left_indent = Inches (0.25)
    p.add_run('PROPOZYCJA PUNKTU DO ROZKAZU REKTORA -KOMENDANTA  WAT').bold = True
    p2 = document.add_paragraph()
    #p2.paragraph_format.left_indent = Inches (0.25)
    p2.add_run('IX. INNE SPRAWY').bold = True

    document.add_paragraph('Stwierdzam, że niżej wymienieni podchorążowie z 2 Batalionu Szkolnego WAT, odbyli przejazd na koszt wojska:', style='List Number')

# dodanie konca dokumentu
def dodaj_stopke(document):

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)

    #p.paragraph_format.left_indent = Inches(0.25)
    p.add_run('Koszty przejazdu opłacić - stanowisko kosztów nr 506 7341 01 UPS 472 15')


    p2 = document.add_paragraph()
    p2.space_before = Pt(12)
    p2.paragraph_format.left_indent = Inches(0.75)
    p2.paragraph_format.first_line_indent = Inches(-0.75)
   # p2.
    p2.add_run('Podstawa: § 2 pkt 2 lit. b, w związku z § 3 rozporządzenia Rady Ministrów z dnia 19 września 2006 r. w sprawie szczególnych uprawnień żołnierzy w czynnej służbie wojskowej do przejazdów na  koszt wojska (Dz. U. z 2014 r., poz. 207).')