import calendar
import datetime
import os
from datetime import timedelta, date

from django.contrib.auth.models import User
from django.http import HttpResponse
from django.utils.dateparse import parse_date
from docx.shared import Inches, Pt, Cm
from docx.enum.table import WD_ROW_HEIGHT
from generator.models import Dane, UserProfile
from docx import Document
from docxtpl import DocxTemplate
from generator.kwotaslownie import kwotaslownie

# ze zmiennej request zbiera się informacje, np kto jest zalogowany

from .maketable import switch_stopien, dodaj_stopke, dodaj_tabele, switch_litery, dodaj_naglowek


def generuj_ext(request, form, generated_doc):
    THIS_FOLDER = os.path.dirname(os.path.abspath(__file__))
    sample_pj = os.path.join(THIS_FOLDER, 'sample_pj.docx')
    sample_ur = os.path.join(THIS_FOLDER, 'sample_ur.docx')



    typ_pociagu = form.cleaned_data.get('typ_pociagu')
    typ_autobusu = form.cleaned_data.get('typ_autobusu')
    temp_typ_srodka = ""
    srodek = ""

    kwota = float(request.POST.get('kwota'))
    if typ_pociagu:
        srodek = "kolejowym w klasie 2, w pociągu "
        for i in typ_pociagu:
            temp_typ_srodka += i + ", "
    elif typ_autobusu:
        srodek = "autobusowym w komunikacji "
        for i in typ_autobusu:
            temp_typ_srodka += i + ", "
    typ_srodka = temp_typ_srodka[:-2]

    if request.POST.get('typ') == 'przepustkę jednorazową':
        doc = DocxTemplate(open(sample_pj, "rb"))
    elif request.POST.get('typ') == 'urlop':
        doc = DocxTemplate(open(sample_ur, "rb"))
    data_wyjazdu = request.POST.get('data_wyjazdu')
    data_powrotu = request.POST.get('data_powrotu')
    miasto = request.POST.get('miasto')
    this_user = User.objects.get(id=request.user.id)
    stopien = this_user.userprofile.stopien
    imie = this_user.userprofile.imie
    nazwisko = this_user.userprofile.nazwisko
    adres = this_user.userprofile.adres
    pluton = this_user.userprofile.pluton
    typ = request.POST.get('typ')
    miesiac = request.POST.get('miesiac')

    # przypisanie id stopnia do stopnia
    wyliczony_stopien = switch_stopien(stopien)

    # liczenie nr rozkazu batalionowego
    wstepna_data = parse_date(request.POST.get('data_wyjazdu'))
    # wstepna_data = parse_date(date(data_wyjazdu.year))
    data_zlozenia = str(date.today()+ timedelta(days=3)) + " r."
    data_rozkazu = '.............'
    nr_rozkazu = '.............'
    if wstepna_data.weekday() == 5 or wstepna_data.weekday() == 4:
        if wstepna_data.weekday() == 5:
            data_rozkazu = wstepna_data - timedelta(days=2)
        else:
            data_rozkazu = wstepna_data - timedelta(days=1)

        # dnia 24.09.2020 był nr rozkazu 76
        # odejmij datę wyjazdu od 24.09, uzyskane dni pocziel na 7 i pomnóż razy 2 - spodziewana liczba wydanych rozkazow
        nr_rozkazu = 0
        # 13.01.2022 był 3 rozkaz wydany
        if(data_rozkazu.year == 2022):
            nr_rozkazu = int(3 + (((data_rozkazu - date(2022, 1, 13)).days + 1) / 7) * 2)
        if(data_rozkazu.year == 2021):
            nr_rozkazu = int(1 + (((data_rozkazu - date(2021, 1, 5)).days + 1) / 7) * 2)
        if(data_rozkazu.year == 2020):
            nr_rozkazu = int(76 + (((data_rozkazu - date(2020, 9, 24)).days + 1) / 7) * 2)

    context = {'stopien':stopien,
               'imie': imie,
               'nazwisko':nazwisko,
               'adres':adres,
               'pluton':pluton,
               'data_przed': parse_date(request.POST.get('data_wyjazdu')) - timedelta(days=1),
               'data_wyjazdu': request.POST.get('data_wyjazdu'),
               'data_powrotu': request.POST.get('data_powrotu'),
               'miesiac': request.POST.get('miesiac'),
               'miejscowosc': request.POST.get('miasto'),
               'kwota': kwota,
               'kwota_slownie': kwotaslownie(kwota, 1),
               'typ': request.POST.get('typ'),
               'typ_srodka': typ_srodka,
               'srodek': srodek,
               'powrot': request.POST.get('tam_z_powrotem'),
               'data_rozkazu': data_rozkazu,
               'nr_rozkazu': nr_rozkazu,
                'data_zlozenia': data_zlozenia,
               }
    nazwisko = nazwisko.upper()
    # tworzenie obiektu bazy danych
    if typ == 'przepustkę jednorazową':
        q = Dane.objects.filter(imie=imie, nazwisko=nazwisko, typ=typ, miesiac=miesiac)
        if q.exists():  # jeśli obiekt istnieje, zaktualizuj jego dane
            dana = Dane.objects.get(imie=imie, nazwisko=nazwisko, typ=typ, miesiac=miesiac)
            dana.data_wyjazdu = data_wyjazdu
            dana.data_powrotu = data_powrotu
            dana.miasto = miasto
            dana.stopien = stopien
            dana.typ = typ
            dana.transport = srodek
            dana.miesiac = miesiac
            dana.stopien_id = wyliczony_stopien
            dana.nr_rozkazu = nr_rozkazu
            dana.save()
        else:
            rekord = Dane(data_wyjazdu=data_wyjazdu, data_powrotu=data_powrotu, miasto=miasto, stopien=stopien,
                          imie=imie, nazwisko=nazwisko, typ=typ, transport=srodek, miesiac=miesiac,
                          stopien_id=wyliczony_stopien, nr_rozkazu=nr_rozkazu)
            rekord.save()
    else:
        q = Dane.objects.filter(imie=imie, nazwisko=nazwisko, typ=typ, data_wyjazdu=data_wyjazdu,
                                data_powrotu=data_powrotu)
        if q.exists():  # jeśli obiekt istnieje, zaktualizuj jego dane
            dana = Dane.objects.get(imie=imie, nazwisko=nazwisko, typ=typ, data_wyjazdu=data_wyjazdu,
                                    data_powrotu=data_powrotu)
            dana.data_wyjazdu = data_wyjazdu
            dana.data_powrotu = data_powrotu
            dana.miasto = miasto
            dana.stopien = stopien
            dana.typ = typ
            dana.transport = srodek
            dana.stopien_id = wyliczony_stopien
            dana.nr_rozkazu = nr_rozkazu

            dana.save()
        else:
            rekord = Dane(data_wyjazdu=data_wyjazdu, data_powrotu=data_powrotu, miasto=miasto, stopien=stopien,
                          imie=imie, nazwisko=nazwisko, typ=typ, transport=srodek, stopien_id=wyliczony_stopien,
                          nr_rozkazu=nr_rozkazu)
            rekord.save()
    doc.render(context)
    doc.save(generated_doc)

def generuj_rozkaz(request):
    query1 = Dane.objects.filter(typ='przepustkę jednorazową', transport__contains='kolejowym w klasie 2').order_by(
        '-stopien_id', 'nazwisko')
    query2 = Dane.objects.filter(typ='urlop', transport__contains='kolejowym w klasie 2').order_by('-stopien_id',
                                                                                                     'nazwisko')
    query3 = Dane.objects.filter(typ='przepustkę jednorazową', transport__contains='autobusowym w').order_by(
        '-stopien_id', 'nazwisko')
    query4 = Dane.objects.filter(typ='urlop', transport__contains='autobusowym w').order_by('-stopien_id',
                                                                                               'nazwisko')

    THIS_FOLDER = os.path.dirname(os.path.abspath(__file__))
    generated_rozkaz = os.path.join(THIS_FOLDER, 'demo.docx')
    document = Document()

    dodaj_naglowek(document)

    licznik = 1

    tables = []

    # tworzenie tabeli osobna funkcja
    if query1:
        p = document.add_paragraph(switch_litery(licznik) + 'na przepustkę jednorazową – środek transportu PKP:')
        # p.paragraph_format.left_indent = Inches (0.25)
        table1 = dodaj_tabele(document, query1)
        tables.append(table1)
        p.paragraph_format.space_before = Pt(12)

    if query2:
        licznik += 1
        p = document.add_paragraph(switch_litery(licznik) + 'na urlop – środek transportu PKP:')
        # p.paragraph_format.left_indent = Inches (0.25)
        table2 = dodaj_tabele(document, query2)
        tables.append(table2)
        p.paragraph_format.space_before = Pt(12)
    if query3:
        licznik += 1
        p = document.add_paragraph(switch_litery(licznik) + 'na przepustkę jednorazową – środek transportu PKS:')
        # p.paragraph_format.left_indent = Inches (0.25)
        table3 = dodaj_tabele(document, query3)
        tables.append(table3)
        p.paragraph_format.space_before = Pt(12)
    if query4:
        licznik += 1
        p = document.add_paragraph(switch_litery(licznik) + 'na urlop – środek transportu PKS:')
        # p.paragraph_format.left_indent = Inches (0.25)
        table4 = dodaj_tabele(document, query4)
        tables.append(table4)
        p.paragraph_format.space_before = Pt(12)

    # ustawianie parametrów dokumentu
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ustawianie szerokosci tabelek
    i = 1
    for tbl in tables:
        for row in tbl.rows:
            j = 1
            row.height_rule = WD_ROW_HEIGHT.EXACTLY
            row.height = Inches(0.20)
            for cell in row.cells:
                if j == 1:  # 1)
                    cell.width = Inches(0.1)
                if j == 2:  # stopien
                    cell.width = Inches(1.25)
                if j == 3:  # imie
                    cell.width = Inches(0.9)
                if j == 4:  # nazwisko
                    cell.width = Inches(1.5)
                if j == 5:
                    cell.width = Inches(2.4)
                if j == 6:
                    cell.width = Inches(0.6)
                if j == 7:
                    cell.width = Inches(1.2)
                j += 1
            i += 1

    dodaj_stopke(document)
    # ustawianie marginesow
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.59)
        section.left_margin = Cm(0.75)
        section.right_margin = Cm(1.32)

    # download
    document.save(generated_rozkaz)
    response = HttpResponse(open(generated_rozkaz, 'rb').read())
    response['Content-Type'] = 'text/plain'
    previous_month = datetime.datetime.today().month - 1
    previous_month_name = calendar.month_name[previous_month]  # ustawienie nazwy pliku z rozkazem
    response['Content-Disposition'] = 'attachment; filename= rozkaz"{}".docx'.format(previous_month_name)

    return response
